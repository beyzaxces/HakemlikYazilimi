from tkinter import *
from tkinter import messagebox
from docx import Document
import os

def bilirkisi_raporu():
    doc = Document()
    doc.add_heading('Bilirkişi Raporu', level=1)
    doc.add_paragraph(f"Uyuşmazlık Konusu: {uyusmazlik_konusu.get()}\n")

    if uyusmazlik_konusu.get() == "Trafik DK":
        doc.add_paragraph(f"Sigortalı araç plakası: {sigortali_plaka_entry.get()}")
        doc.add_paragraph(f"Başvurucuya ait araç plakası: {basvuran_plaka_entry.get()}")
        doc.add_paragraph(f"Kaza tarihi: {kaza_tarihi_entry.get()}\n")
        doc.add_paragraph("Sigorta Tahkim Komisyon Başkanlığı tarafından Heyetimize intikal ettirilen uyuşmazlık, "
                          f"davalı sigorta kuruluşuna zorunlu mali sorumluluk sigortalı \"{sigortali_plaka_entry.get()}\" "
                          f"plakalı araç ile başvurucuya ait \"{basvuran_plaka_entry.get()}\" plakalı araç arasında "
                          f"\"{kaza_tarihi_entry.get()}\" tarihinde gerçekleşen trafik kazası sonucu başvurucuya ait "
                          f"araçta oluşan değer kaybı tutarının tazmini amacıyla Komisyona yapılan başvuruda "
                          f"Uyuşmazlık Hakemince verilen Karara davalı sigorta kuruluşunun itirazlarına ilişkindir.")

    doc.add_paragraph(f"\n\nİnceleme süresi: {intikal_sureci.get()}\n")

    if intikal_sureci.get() == "Doğrudan hüküm":
        doc.add_paragraph("\nUyuşmazlık Hakemi Kararına davalı sigorta kuruluşu vekilinin itirazları Komisyon nezdinde "
                          "öncelikle İtiraz Yetkilisi tarafından incelenmiş ve itirazın süresinde ve usulüne uygun "
                          "şekilde yapıldığı tespit edildiğinden, davalının itirazlarının değerlendirilmesi ve "
                          "uyuşmazlığın çözümü için dosya İtiraz Hakem Heyetimize intikal ettirilmiştir.\n\nHeyetimizce "
                          "yapılan incelemede, dosyada mevcut bilgi ve delillerin uyuşmazlık ve itirazlar konusunda "
                          "bir kanaate ulaşabilmek için yeterli olduğu kanaatine ulaşılmış ve doğrudan hüküm "
                          "kurulması yoluna gidilmiştir.\n\nHeyetimizce davalı vekilinin itirazlarına ilişkin olarak "
                          "karara varılmış ve yargılamaya son verilmiştir.")

    doc.add_heading('SİGORTA HAKEMİ VEYA HAKEM HEYETİNCE VERİLEN HÜKÜM', level=1)
    doc.add_paragraph(f"Uyuşmazlık Hakemi tarafından verilen hüküm: {hukum.get()}")

    if hukum.get() == "Başvurunun kabulü":
        if not deger_kaybi_entry.get().strip() or not vekalet_ucreti_entry.get().strip():
            messagebox.showwarning("Eksik Bilgi", "Lütfen değer kaybı ve vekalet ücreti bilgilerini girin.")
            return

        try:
            deger_kaybi = float(deger_kaybi_entry.get())
            vekalet_ucreti = float(vekalet_ucreti_entry.get())
        except ValueError:
            messagebox.showwarning("Geçersiz Değer", "Değer kaybı ve vekalet ücreti sayı olmalıdır.")
            return

        karar_metni = f"\nUyuşmazlık Hakemi tarafından {deger_kaybi} TL değer kaybı bedeli, yargılama giderleri ve {vekalet_ucreti} TL vekalet ücretinin davalı sigorta kuruluşu tarafından başvurucuya ödenmesine karar verilmiştir."
        doc.add_paragraph(karar_metni)

    # Word dosyasını kaydet
    dosya_yolu = os.path.join("C:\\Users\\Beyza\\Desktop", "BilirkisiDeneme_Raporu.docx")
    doc.save(dosya_yolu)

    # Bilgilendirme mesajı
    messagebox.showinfo("İşlem Tamamlandı", f"Bilirkişi raporu oluşturuldu: {dosya_yolu}")

def uyusmazlik_goster():
    # Uyuşmazlık Konusu seçildiğinde ilgili inputları göster veya gizle
    if uyusmazlik_konusu.get() == "Trafik DK":
        sigortali_plaka_label.grid()
        sigortali_plaka_entry.grid()
        basvuran_plaka_label.grid()
        basvuran_plaka_entry.grid()
        kaza_tarihi_label.grid()
        kaza_tarihi_entry.grid()
    else:
        sigortali_plaka_label.grid_remove()
        sigortali_plaka_entry.grid_remove()
        basvuran_plaka_label.grid_remove()
        basvuran_plaka_entry.grid_remove()
        kaza_tarihi_label.grid_remove()
        kaza_tarihi_entry.grid_remove()

def hukum_secildi(*args):
    if hukum.get() == "Başvurunun kabulü":
        deger_kaybi_label.grid(row=19, column=0, sticky=W, padx=10, pady=5)
        deger_kaybi_entry.grid(row=19, column=1, padx=10, pady=5)
        vekalet_ucreti_label.grid(row=20, column=0, sticky=W, padx=10, pady=5)
        vekalet_ucreti_entry.grid(row=20, column=1, padx=10, pady=5)
    else:
        deger_kaybi_label.grid_remove()
        deger_kaybi_entry.grid_remove()
        vekalet_ucreti_label.grid_remove()
        vekalet_ucreti_entry.grid_remove()

# Ana pencereyi oluştur
root = Tk()
root.title("Bilirkişi Raporu ve Sigorta Hakemi veya Hakem Heyeti Hükmü")

# Uyuşmazlık Konusu
uyusmazlik_konusu = StringVar()
uyusmazlik_konusu.set("Trafik DK")
Label(root, text="Uyuşmazlığın Konusu:").grid(row=0, column=0, sticky=W, padx=10, pady=5)
uyusmazlik_secenekleri = [
    "Trafik DK", "Trafik Hasar", "Trafik DK ve Hasar", "Kasko",
    "Trafik SİGT", "Trafik DYKT", "İMM", "Zorunlu Deprem", "Diğer"
]
for idx, option in enumerate(uyusmazlik_secenekleri):
    Radiobutton(root, text=option, variable=uyusmazlik_konusu, value=option, command=uyusmazlik_goster).grid(
        row=idx + 1, column=0, sticky=W, padx=10, pady=5)

# Trafik DK seçeneği için input alanları
sigortali_plaka_label = Label(root, text="Sigortalı araç plakası:")
sigortali_plaka_entry = Entry(root)
sigortali_plaka_label.grid(row=10, column=0, sticky=W, padx=10, pady=5)
sigortali_plaka_entry.grid(row=10, column=1, padx=10, pady=5)

basvuran_plaka_label = Label(root, text="Başvurucuya ait araç plakası:")
basvuran_plaka_entry = Entry(root)
basvuran_plaka_label.grid(row=11, column=0, sticky=W, padx=10, pady=5)
basvuran_plaka_entry.grid(row=11, column=1, padx=10, pady=5)

kaza_tarihi_label = Label(root, text="Kaza tarihi (DD/MM/YYYY):")
kaza_tarihi_entry = Entry(root)
kaza_tarihi_label.grid(row=12, column=0, sticky=W, padx=10, pady=5)
kaza_tarihi_entry.grid(row=12, column=1, padx=10, pady=5)

# Uyuşmazlık Hakemi tarafından verilen hüküm
Label(root, text="Uyuşmazlık Hakemi tarafından verilen hüküm:").grid(row=13, column=0, sticky=W, padx=10, pady=5)
hukum = StringVar()
hukum.set("Başvurunun kabulü")
for idx, option in enumerate(["Başvurunun kabulü", "Başvurunun kısmen kabulü", "Başvurunun reddi", "Başvurunun usulden reddi"]):
    Radiobutton(root, text=option, variable=hukum, value=option).grid(row=14 + idx, column=0, sticky=W, padx=10, pady=5)

# Değer kaybı ve vekalet ücreti için input alanları
deger_kaybi_label = Label(root, text="Değer Kaybı (TL):")
deger_kaybi_entry = Entry(root)
vekalet_ucreti_label = Label(root, text="Vekalet Ücreti (TL):")
vekalet_ucreti_entry = Entry(root)

# Gönder butonu
Button(root, text='Gönder', command=bilirkisi_raporu).grid(row=21, column=0, padx=10, pady=10)

# Hüküm seçildiğinde fonksiyonu bağlama
hukum.trace_add("write", hukum_secildi)

# Pencereyi aç
root.mainloop()
