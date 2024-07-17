from tkinter import *
from tkinter import messagebox
from docx import Document
import os

def bilirkisi_raporu():
    doc = Document()
    doc.add_heading('BİLİRKİŞİ RAPORU', level=1)
    doc.add_heading('1. BAŞVURUYA KONU UYUŞMAZLIK VE İTİRAZ HAKEM HEYETİNE İNTİKALİ ', level=2)
    doc.add_heading(f"1.1 Uyuşmazlık Konusu: {uyusmazlik_konusu.get()}\n", level=3)

    if uyusmazlik_konusu.get() == "Trafik DK":
        doc.add_paragraph(f"Sigortalı araç plakası: {sigortali_plaka_entry.get()}")
        doc.add_paragraph(f"Başvurucuya ait araç plakası: {basvuran_plaka_entry.get()}")
        doc.add_paragraph(f"Kaza tarihi: {kaza_tarihi_entry.get()}\n")
        doc.add_paragraph("Sigorta Tahkim Komisyon Başkanlığı tarafından Heyetimize intikal ettirilen uyuşmazlık, "
                          f"davalı sigorta kuruluşuna zorunlu mali sorumluluk sigortalı \"{sigortali_plaka_entry.get()}\" "
                          f"plakalı araç ile başvurucuya ait \"{basvuran_plaka_entry.get()}\" plakalı araç arasında "
                          f"\"{kaza_tarihi_entry.get()}\" tarihinde gerçekleşen trafik kazası sonucu başvurucuya ait "
                          f"araçta oluşan değer kaybı tutarının tazmini amacıyla Komisyona yapılan başvuruda "
                          f"Uyuşmazlık Hakemince verilen Karara davalı sigorta kuruluşunun itirazlarına ilişkindir.")

    doc.add_heading(f"\n1.2 İnceleme süresi: {intikal_sureci.get()}\n", level=3)

    if intikal_sureci.get() == "Doğrudan hüküm":
        doc.add_paragraph("\nUyuşmazlık Hakemi Kararına davalı sigorta kuruluşu vekilinin itirazları Komisyon nezdinde "
                          "öncelikle İtiraz Yetkilisi tarafından incelenmiş ve itirazın süresinde ve usulüne uygun "
                          "şekilde yapıldığı tespit edildiğinden, davalının itirazlarının değerlendirilmesi ve "
                          "uyuşmazlığın çözümü için dosya İtiraz Hakem Heyetimize intikal ettirilmiştir.\n\nHeyetimizce "
                          "yapılan incelemede, dosyada mevcut bilgi ve delillerin uyuşmazlık ve itirazlar konusunda "
                          "bir kanaate ulaşabilmek için yeterli olduğu kanaatine ulaşılmış ve doğrudan hüküm "
                          "kurulması yoluna gidilmiştir.\n\nHeyetimizce davalı vekilinin itirazlarına ilişkin olarak "
                          "karara varılmış ve yargılamaya son verilmiştir.")

    doc.add_heading('2. SİGORTA HAKEMİ VEYA HAKEM HEYETİNCE VERİLEN HÜKÜM', level=2)
    doc.add_heading(f"Uyuşmazlık Hakemi tarafından verilen hüküm: {hukum.get()}", level=3)

    if hukum.get() == "Başvurunun kabulü":
        if not deger_kaybi_entry.get().strip() or not vekalet_ucreti_entry.get().strip():
            messagebox.showwarning("Eksik Bilgi", "Lütfen değer kaybı ve vekalet ücreti bilgilerini girin.")
            return

        try:
            deger_kaybi = float(deger_kaybi_entry.get())
            vekalet_ucreti = float(vekalet_ucreti_entry.get())
        except ValueError:
            messagebox.showwarning("Geçersiz Değer", "Değer kaybı ve vekalet ücreti sayı olmalıdır.")
            return

        karar_metni = f"\nUyuşmazlık Hakemi tarafından {deger_kaybi} TL değer kaybı bedeli, yargılama giderleri ve {vekalet_ucreti} TL vekalet ücretinin davalı sigorta kuruluşu tarafından başvurucuya ödenmesine karar verilmiştir."
        doc.add_paragraph(karar_metni)

    # Word dosyasını kaydet
    dosya_yolu = os.path.join("C:\\Users\\Beyza\\Desktop", "Bilirkisi_Raporu.docx")
    doc.save(dosya_yolu)

    # Bilgilendirme mesajı
    messagebox.showinfo("İşlem Tamamlandı", f"Bilirkişi raporu oluşturuldu: {dosya_yolu}")


# Ana pencereyi oluştur
root = Tk()
root.title("Bilirkişi Raporu")
root.geometry("700x800")

# Uyuşmazlık Konusu
uyusmazlik_konusu = StringVar()
uyusmazlik_konusu.set("Trafik DK")
Label(root, text="Uyuşmazlığın Konusu:").grid(row=0, column=0, sticky=W, padx=10, pady=5)
uyusmazlik_secenekleri = [
    "Trafik DK", "Trafik Hasar", "Trafik DK ve Hasar", "Kasko",
    "Trafik SİGT", "Trafik DYKT", "İMM", "Zorunlu Deprem", "Diğer"
]
for idx, option in enumerate(uyusmazlik_secenekleri):
    Radiobutton(root, text=option, variable=uyusmazlik_konusu, value=option).grid(
        row=idx + 1, column=0, sticky=W, padx=10, pady=5)

# Trafik DK seçeneği için input alanları
sigortali_plaka_label = Label(root, text="Sigortalı araç plakası:")
sigortali_plaka_entry = Entry(root)
sigortali_plaka_label.grid(row=1, column=3, sticky=W, padx=10, pady=5)
sigortali_plaka_entry.grid(row=1, column=4, padx=10, pady=5)

basvuran_plaka_label = Label(root, text="Başvurucuya ait araç plakası:")
basvuran_plaka_entry = Entry(root)
basvuran_plaka_label.grid(row=2, column=3, sticky=W, padx=10, pady=5)
basvuran_plaka_entry.grid(row=2, column=4, padx=10, pady=5)

kaza_tarihi_label = Label(root, text="Kaza tarihi (DD/MM/YYYY):")
kaza_tarihi_entry = Entry(root)
kaza_tarihi_label.grid(row=3, column=3, sticky=W, padx=10, pady=5)
kaza_tarihi_entry.grid(row=3, column=4, padx=10, pady=5)

# İnceleme Süresi
Label(root, text="İnceleme Süresi:").grid(row=10, column=0, sticky=W, padx=10, pady=5)
intikal_sureci = StringVar()
intikal_sureci.set("Doğrudan hüküm")
intikal_secenekleri = ["Doğrudan hüküm", "Taraflardan belge isteme sonrası hüküm", "Bilirkişi tayini sonrası hüküm",
                       "Yeni Rapor istenilmesi sonrası hüküm", "Yeni Rapor istenilmesi ve BK tayini sonrası hüküm"]
intikal_sureci_radiobuttons = [Radiobutton(root, text=option, variable=intikal_sureci, value=option) for option in
                               intikal_secenekleri]
for idx, option in enumerate(intikal_secenekleri):
    intikal_sureci_radiobuttons[idx].grid(row=11 + idx, column=0, sticky=W)

# Uyuşmazlık Hakemi tarafından verilen hüküm
Label(root, text="Uyuşmazlık Hakemi tarafından verilen hüküm:").grid(row=16, column=0, sticky=W, padx=10, pady=5)
hukum = StringVar()
hukum.set("Başvurunun kabulü")
hukum_secenekleri = [
    "Başvurunun kabulü", "Başvurunun kısmen kabulü",
    "Başvurunun reddi", "Başvurunun usulden reddi"
]
for idx, option in enumerate(hukum_secenekleri):
    Radiobutton(root, text=option, variable=hukum, value=option).grid(row=17 + idx, column=0, sticky=W, padx=10, pady=5)

# Değer kaybı ve vekalet ücreti için input alanları
deger_kaybi_label = Label(root, text="Değer Kaybı (TL):")
deger_kaybi_entry = Entry(root)
vekalet_ucreti_label = Label(root, text="Vekalet Ücreti (TL):")
vekalet_ucreti_entry = Entry(root)
vekalet_ucreti_label.grid(row=17, column=3, sticky=W, padx=10, pady=5)
vekalet_ucreti_entry.grid(row=17, column=4, padx=10, pady=5)
deger_kaybi_label.grid(row=18, column=3, sticky=W, padx=10, pady=5)
deger_kaybi_entry.grid(row=18, column=4, sticky=W, padx=10, pady=5)



# Gönder butonu
Button(root, text='Gönder', command=bilirkisi_raporu).grid(row=26, column=0, padx=10, pady=10)

# Pencereyi aç
root.mainloop()
