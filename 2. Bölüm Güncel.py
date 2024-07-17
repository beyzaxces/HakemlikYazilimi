from tkinter import *
from tkinter import messagebox
from docx import Document

# Ana pencereyi oluştur
root = Tk()
root.title("Sigorta Hakemi veya Hakem Heyeti Hükmü")

# Uyuşmazlık Hakemi tarafından verilen hüküm
Label(root, text="Uyuşmazlık Hakemi tarafından verilen hüküm:").grid(row=0, column=0, sticky=W)

hukum = StringVar()
hukum.set("Başvurunun kabulü")
hukum_secenekleri = ["Başvurunun kabulü", "Başvurunun kısmen kabulü", "Başvurunun reddi", "Başvurunun usulden reddi"]
for idx, option in enumerate(hukum_secenekleri):
    Radiobutton(root, text=option, variable=hukum, value=option).grid(row=idx + 1, column=0, sticky=W)

# Değer kaybı ve vekalet ücreti inputları
deger_kaybi_label = Label(root, text="Değer kaybı (TL):")
deger_kaybi_entry = Entry(root)

vekalet_ucreti_label = Label(root, text="Vekalet ücreti (TL):")
vekalet_ucreti_entry = Entry(root)

# Gönder butonu
def sigorta_hukumu():
    # Word belgesi oluştur
    doc = Document()

    # Başlık ve başlangıç metni ekle
    doc.add_paragraph('Sigorta Hakemi Veya Hakem Heyetince Verilen Hüküm\n')
    doc.add_paragraph(f"Uyuşmazlık Hakemi tarafından verilen hüküm: {hukum.get()}")

    if hukum.get() == "Başvurunun kabulü":
        if not deger_kaybi_entry.get().strip() or not vekalet_ucreti_entry.get().strip():
            messagebox.showwarning("Eksik Bilgi", "Lütfen değer kaybı ve vekalet ücreti bilgilerini girin.")
            return

        try:
            deger_kaybi = float(deger_kaybi_entry.get())
            vekalet_ucreti = float(vekalet_ucreti_entry.get())
        except ValueError:
            messagebox.showwarning("Geçersiz Değer", "Değer kaybı ve vekalet ücreti sayı olmalıdır.")
            return

        karar_metni = f"\nUyuşmazlık Hakemi tarafından {deger_kaybi} TL değer kaybı bedeli, yargılama giderleri ve {vekalet_ucreti} TL vekalet ücretinin davalı sigorta kuruluşu tarafından başvurucuya ödenmesine karar verilmiştir."
        doc.add_paragraph(karar_metni)

    # Word dosyasını kaydet
    doc.save(r"C:\Users\Beyza\Desktop\Bilirkisi2.docx")
    messagebox.showinfo("İşlem Tamamlandı", "Bilirkisi2.docx dosyası başarıyla kaydedildi.")

gonder_butonu = Button(root, text='Gönder', command=sigorta_hukumu)
gonder_butonu.grid(row=7, column=1, sticky=W, pady=4)


# Hüküm seçimini takip etmek için event binding
def hukum_secildi(*args):
    if hukum.get() == "Başvurunun kabulü":
        deger_kaybi_label.grid(row=5, column=0, sticky=W)
        deger_kaybi_entry.grid(row=5, column=1)

        vekalet_ucreti_label.grid(row=6, column=0, sticky=W)
        vekalet_ucreti_entry.grid(row=6, column=1)
    else:
        deger_kaybi_label.grid_remove()
        deger_kaybi_entry.grid_remove()

        vekalet_ucreti_label.grid_remove()
        vekalet_ucreti_entry.grid_remove()

# Hüküm seçimi için trace ekleme
hukum.trace_add('write', hukum_secildi)

root.mainloop()
